from __future__ import annotations

import io


def extract_text(uploaded_file) -> str:
    """Return clean plain text from a Streamlit UploadedFile (.pdf or .docx)."""
    name = uploaded_file.name.lower()
    raw: bytes = uploaded_file.read()

    if name.endswith(".pdf"):
        return _from_pdf(raw)
    if name.endswith(".docx"):
        return _from_docx(raw)
    raise ValueError(
        f"Unsupported file type '{uploaded_file.name}'. Please upload a .pdf or .docx file."
    )


def _from_pdf(data: bytes) -> str:
    import pdfplumber

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]

    text = "\n\n".join(p.strip() for p in pages if p.strip())
    if not text:
        raise ValueError(
            "No text could be extracted from the PDF. "
            "It may be a scanned / image-based document — please paste the text manually."
        )
    return text


def _from_docx(data: bytes) -> str:
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(data))

    # Collect paragraphs + table cells so nothing is missed
    lines: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in lines:
                    lines.append(t)

    if not lines:
        raise ValueError(
            "The DOCX file appears to be empty or has no readable text."
        )
    return "\n\n".join(lines)
